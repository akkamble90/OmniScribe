import pytesseract
from PIL import Image
from docx import Document as DocxDocument
from langchain_community.document_loaders import UnstructuredURLLoader, PyPDFLoader, WebBaseLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

def process_any_format(file_path: str, file_type: str):
    """Processes PDF, Word, Images, and WebLinks into Split LangChain Documents."""
    docs = []
    
    try:
        if file_type == "pdf":
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            
        elif file_type == "docx":
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            docs = [Document(page_content=full_text, metadata={"source": file_path})]
            
        elif file_type in ["png", "jpg", "jpeg", "bmp"]:
            # OCR for images
            text = pytesseract.image_to_string(Image.open(file_path))
            docs = [Document(page_content=text, metadata={"source": file_path})]
            
        elif file_type == "url":
            # --- IMPROVED HYBRID WEB SCRAPING ---
            try:
                # 1. Try WebBaseLoader first (Faster, lighter, more reliable for basic HTML)
                loader = WebBaseLoader(file_path)
                loader.verify_ssl = False  # Bypasses SSL certificate errors
                # Mimic a real browser to prevent being blocked
                loader.requests_kwargs = {
                    "headers": {
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
                    }
                }
                docs = loader.load()
            except Exception:
                # 2. Fallback to Unstructured if WebBaseLoader fails
                loader = UnstructuredURLLoader(
                    urls=[file_path], 
                    ssl_verify=False,
                    headers={
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
                    }
                )
                docs = loader.load()

        # --- VALIDATION: Ensure we actually got text ---
        docs = [d for d in docs if d.page_content.strip()]
        
        if not docs:
            print(f"Warning: No text content extracted from {file_type}")
            return []

        # --- TEXT SPLITTING ---
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
        split_docs = text_splitter.split_documents(docs)
        print(f"Successfully processed {file_type}: Created {len(split_docs)} chunks.")
        return split_docs

    except Exception as e:
        print(f"Error processing {file_type}: {str(e)}")
        return []